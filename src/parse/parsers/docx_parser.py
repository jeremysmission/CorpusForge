"""
DOCX parser -- reads modern Microsoft Word (.docx) files.

Plain English: opens a Word document and pulls out the body paragraphs
as clean text. Uses the ``python-docx`` library. If the file is corrupt
or the library isn't available, the parser logs the issue and returns
an empty document -- it never crashes the Forge pipeline.
"""

from __future__ import annotations

import logging
from pathlib import Path

from src.parse.parsers.txt_parser import ParsedDocument

logger = logging.getLogger(__name__)


class DocxParser:
    """Extract body text from modern Word .docx documents."""

    def parse(self, file_path: Path) -> ParsedDocument:
        """Open a .docx file and return its paragraphs as clean text."""
        path = Path(file_path)
        try:
            from docx import Document
            doc = Document(str(path))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            text = "\n\n".join(paragraphs)
        except Exception as e:
            logger.error("DOCX parse failed for %s: %s", path.name, e)
            text = ""

        return ParsedDocument(
            source_path=str(path),
            text=text,
            parse_quality=1.0 if text.strip() else 0.0,
            file_ext=".docx",
            file_size=path.stat().st_size,
        )
